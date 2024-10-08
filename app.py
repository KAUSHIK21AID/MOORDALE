from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify
import pandas as pd
from io import BytesIO
import os
from docx import Document
import google.generativeai as genai
from scholarly import scholarly
import concurrent.futures
import matplotlib.pyplot as plt
import seaborn as sns
from flask_cors import CORS

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
app.secret_key = 'supersecretkey'  # Ensure you have a secret key for sessions and flash messages

# Initialize Google Generative AI
genai.configure(api_key="AIzaSyDJkX7rIew0l3siFeVYZAIh2xNVGsavmRk") # Replace with your actual API key
model = genai.GenerativeModel('gemini-1.5-flash-latest')

# Global variables
processed_df = None  # Placeholder for DataFrame
results_html = None  # Placeholder for HTML table

# Function to analyze author data
def d_analysis(author_name):
  search_query = scholarly.search_author(author_name)
  first_author_result = next(search_query)
  author = scholarly.fill(first_author_result)
  years = []
  for i in author['publications']:
    years.append(i['bib'].get('pub_year'))
  yy = [year for year in years if year is not None]

  hindex = author['hindex']
  iindex = author['i10index']
  year_count = {}


  for year in yy:
      if year in year_count:
          year_count[year] += 1
      else:
          year_count[year] = 1

  ordered_year_count = dict(sorted(year_count.items()))
  jif=[]
  yearlist=[]
  for key, val in author['cites_per_year'].items():
    sum = 0
    for kk, val in ordered_year_count.items():
      if(int(kk)==key-1 or int(kk)==key-2):
        sum = sum + val
      elif(int(kk)>key-1):
        break
    if(sum==0):
      continue
    jif.append(val/sum)
    yearlist.append(key)

    d_a = pd.DataFrame()
    d_a['Author'] = [author_name]*len(jif)
    d_a['Year'] = yearlist
    d_a['JIF'] = jif
    d_a['H-index'] = [hindex]*len(jif)
    d_a['I-index'] = [iindex]*len(jif)


  return d_a

# Function to retrieve scholarly information
def retrieve_stuffs(author_name, institution_name):
    try:
        search_query = f"{author_name} {institution_name}"
        search_result = scholarly.search_author(search_query)
        first_author_result = next(search_result)
        author = scholarly.fill(first_author_result)

        titles = [pub['bib']['title'] for pub in author['publications']]
        years = [pub['bib'].get('pub_year') for pub in author['publications']]
        citation = [pub['bib'].get('citation', '').lower() for pub in author['publications']]

        df = pd.DataFrame({
            'Author': [author_name] * len(author['publications']),
            'Title': titles,
            'Publication Year': years,
            'Citation': citation,
            'Institution_Name': institution_name
        })

        return df
    except StopIteration:
        print(f"Author not found: {author_name}")
        return pd.DataFrame()
    except Exception as e:
        print(f"Error processing {author_name}: {e}")
        return pd.DataFrame()

def generate_author_summary(df, author):
    filtered_df = df[df['Author'] == author]
    if filtered_df.empty:
        return f"No publications found for author '{author}'"

    titles = filtered_df['Title'].tolist()
    citations = filtered_df['Citation'].tolist()
    intro_phrases = [
        "One of the key works was",
        "Another notable publication was",
        "Among the significant contributions was",
        "An important study was",
        "A remarkable work was",
        "Additionally, there was",
        "Another critical work was"
    ]
    summary = f"{author} made significant contributions to their field with several impactful publications. "
    for i, (title, citation) in enumerate(zip(titles, citations)):
        phrase = intro_phrases[i % len(intro_phrases)]
        summary += f"{phrase} '{title}', which appeared in {citation}. "
    summary += f"These publications underscore {author}'s commitment to advancing research in their domain, particularly in areas such as {', '.join([title.split(':')[0].lower() for title in titles[:2]])}, and other related fields."

    SYSTEM_PROMPT = "Your name is Summarize AI. Your task is to Summarize the context."

    chat = model.start_chat(history=[{"role": "model", "parts": [SYSTEM_PROMPT]}])
    response = chat.send_message(f"""
User Prompt: Summarize the following text.
\n\n
Here is the Text: {summary}
\n\n
Instruction:
Generate a concise summary of the provided text.
""", stream=True)

    final_summary = ""
    for chunk in response:
        final_summary += chunk.text

    return final_summary

def generate_word_doc(summary_text):
    doc = Document()
    doc.add_heading('Customized Summary', level=1)
    doc.add_paragraph(summary_text)
    return doc

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)

    file = request.files['file']

    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)

    if file:
        try:
            df = pd.read_csv(file)
            combined_df = pd.DataFrame()

            institution_name = df['Institution_Name'].iloc[0] if 'Institution_Name' in df.columns else ''
            author_list = df['Author'].tolist()

            with concurrent.futures.ThreadPoolExecutor() as executor:
                results = list(executor.map(lambda name: retrieve_stuffs(name, institution_name), author_list))

            for result in results:
                combined_df = pd.concat([combined_df, result], ignore_index=True)

            global processed_df, results_html
            processed_df = combined_df
            results_html = combined_df.to_html(classes='table dataTable', index=False)

            return redirect(url_for('results'))

        except Exception as e:
            flash(f'An error occurred while processing the file: {str(e)}')
            return redirect(url_for('index'))

@app.route('/results')
def results():
    if processed_df is None:
        flash("No results to display. Please upload and process a CSV first.")
        return redirect(url_for('index'))
    
    global results_html
    results_html = processed_df.to_html(classes='table dataTable', index=False)
    return render_template('results.html', results=results_html, authors=processed_df['Author'].unique().tolist())

@app.route('/summary', methods=['POST'])
def summary():
    author = request.form.get('author')
    if author:
        summary_text = generate_author_summary(processed_df, author)
        return jsonify({'summary': summary_text})
    else:
        return jsonify({'summary': 'No author selected.'})

@app.route('/download_summary', methods=['GET'])
def download_summary():
    author = request.args.get('author')
    if author:
        summary_text = generate_author_summary(processed_df, author)
        summary_doc = generate_word_doc(summary_text)
        summary_buffer = BytesIO()
        summary_doc.save(summary_buffer)
        summary_buffer.seek(0)
        return send_file(summary_buffer, as_attachment=True, download_name=f"{author}_summary.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        flash("No author selected.")
        return redirect(url_for('results'))

@app.route('/download')
def download():
    if processed_df is None:
        flash("No data to download. Please upload and process a CSV first.")
        return redirect(url_for('index'))

    csv_buffer = BytesIO()
    processed_df.to_csv(csv_buffer, index=False)
    csv_buffer.seek(0)
    
    return send_file(csv_buffer, as_attachment=True, download_name='results.csv', mimetype='text/csv')

@app.route('/view_analysis', methods=['POST'])
def view_analysis():
    global processed_df
    
    if processed_df is None:
        return jsonify({'analysis': 'No data available for analysis.'})

    # Get unique authors from the processed DataFrame
    authors = processed_df['Author'].unique().tolist()

    analysis_dfs = [d_analysis(author) for author in authors]
    combined_analysis_df = pd.concat(analysis_dfs, ignore_index=True)

    return jsonify({
        'analysis': combined_analysis_df.to_html(classes='table dataTable', index=False),
        'chart_data': combined_analysis_df['Author'].value_counts().to_dict()
    })

if __name__ == '__main__':
    app.run(debug=True)
